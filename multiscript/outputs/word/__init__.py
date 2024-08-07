
from enum import Enum, auto
import logging

import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from multiscript.outputs.base import OutputVersionConfig, OutputPlanConfig, OutputPlanRun, OutputBibleStreamHandler
from multiscript.outputs.tagged import TaggedOutput, TaggedDocCursor, Tags
from multiscript.outputs.word.version_config_word_panel import WordVersionConfigPanel
from multiscript.outputs.word.plan_config_word_panel import WordPlanConfigPanel
from multiscript.plan.symbols import column_symbols


_logger = logging.getLogger(__name__)


class Styles(Enum):
    PARAGRAPH               =   "MSC_Paragraph"             # Base style for Bible text.
    PARAGRAPH_COL           =   "MSC_Paragraph_{0}"         # Style for Bible text for a particular column.
    PASSAGE                 =   "MSC_Passage"               # Style for a single passage reference.
    PASSAGE_GROUP           =   "MSC_Passage_Group"         # Style for a group of passage references.
    TEXT_TABLE_HORIZ        =   "MSC_Text_Table_Horiz"      # Style for horizontal table of Bible texts.
    VERSION_NAME            =   "MSC_Version_Name"          # Style for the Bible version's name.
    JOIN                    =   "MSC_Join"                  # Style for the join indicator between group passages.
    JOIN_COL                =   "MSC_Join_{0}"              # Style for the join indicator for a particular column.
    COPYRIGHT_TABLE_HORIZ   =   "MSC_Copyright_Table_Horiz" # Style for horizontal table of copyright texts.
    COPYRIGHT               =   "MSC_Copyright"             # Style for the copyright text.


class WordOutput(TaggedOutput):
    
    ID = "word"

    def __init__(self, plugin):
        super().__init__(plugin)
        self.id = WordOutput.ID
        self.name = "MS Word"
        self.output_file_ext = ".docx"
        self._accepted_template_exts.append(".docx")

    #
    # Implementation of abstract methods from BibleOutput
    #
    
    def new_output_version_config(self):
        return WordVersionConfig(self)

    def new_output_app_config(self):
        return None

    def new_output_plan_config(self):
        return WordPlanConfig(self)

    def new_output_plan_run(self, plan):
        return WordPlanRun(plan)

    def setup(self, runner):
        '''Overriden from TaggedOutput.setup(). Called prior to looping through the version
        combos.

        We use this method to set some defaults for the run.
        '''
        super().setup(runner)
        runner.output_runs[self.long_id].text_join = runner.plan.config.outputs[self.long_id].join_passage_text
    
    #
    # Implementation of abstract methods from FileSetOutput
    #

    def load_document(self, runner, version_combo, template_path):
        document = docx.Document(template_path)

        # Apply formatting to styles
        for combo_element in version_combo:
            column_index = combo_element.version_column.symbol_index
            bible_version = combo_element.version

            # Column paragraph styles
            col_para_style_name = Styles.PARAGRAPH_COL.value.format(column_symbols[column_index])
            col_para_style = get_style(document, col_para_style_name)
            if col_para_style is None:
                # Style is missing, so create it.
                col_para_style = document.styles.add_style(col_para_style_name, WD_STYLE_TYPE.PARAGRAPH)
                base_style = get_style(document, Styles.PARAGRAPH.value)
                if base_style is None:
                    base_style = get_style(document, "Normal")
                col_para_style.base_style = base_style
            if bible_version is not None and not runner.plan.config.outputs[self.long_id].apply_direct_formatting:
                set_run_formatting(col_para_style, bible_version.font_family,
                                    bible_version.output_config[self.long_id].font_size, bible_version.is_rtl)
                set_para_formatting(col_para_style, bible_version.is_rtl)

            # Column text-join styles
            col_text_join_style_name = Styles.JOIN_COL.value.format(column_symbols[column_index])
            col_text_join_style = get_style(document, col_text_join_style_name)
            if col_text_join_style is None:
                # Style is missing, so create it.
                col_text_join_style = document.styles.add_style(col_text_join_style_name, WD_STYLE_TYPE.PARAGRAPH)
                base_style = get_style(document, Styles.JOIN.value)
                if base_style is None:
                    base_style = get_style(document, "Normal")
                col_text_join_style.base_style = base_style
            if bible_version is not None and not runner.plan.config.outputs[self.long_id].apply_direct_formatting:
                set_run_formatting(col_text_join_style, bible_version.font_family,
                                    bible_version.output_config[self.long_id].font_size, bible_version.is_rtl)
                set_para_formatting(col_text_join_style, bible_version.is_rtl)

        return document
    
    def save_document(self, runner, version_combo, document, filepath):
        document.save(filepath)

    #
    # Implementation of abstract methods from TaggedOutput
    #    

    def replace_tag_directly(self, document, tag_text, replacement_text):
        '''Searches document for tag_text, and replaces it with replacement_text.
        Tries to leave surrounding text in the paragraph untouched, though unfortunately
        sometimes the formatting is lost.
        If tag_text is not found, the document is not modified.
        '''
        # Collect all the paragraphs from the document body plus any in table cells
        all_paragraphs = document.paragraphs.copy()
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_paragraphs += cell.paragraphs
                    
        for paragraph in all_paragraphs:
            if tag_text in paragraph.text:
                for run in paragraph.runs:
                    if tag_text in run.text:
                        text = run.text.split(tag_text)
                        run.text = text[0] + replacement_text + text[1]
                        return
                # At this point, search_text is in paragraph, but not in a single run
                # The search_text will lose its formatting.
                #
                # TODO: Can we avoid losing the formatting?
                #
                text = paragraph.text.split(tag_text)
                paragraph.text = text[0] + replacement_text + text[1]                                

    def replace_tag_with_cursor(self, document, tag_text):
        '''Searches document for tag_text, removes it and returns a subclass of TaggedDocCursor
        pointing to the paragraph. If not tag_text is not found, returns None.

        NOTE: At this time we are unable to preserve the remaining text in the paragraph.
              So any text surrounding this tag in the same paragraph will be cleared.
        '''
        cursor = None

        # Collect all the paragraphs from the doc body plus any in table cells
        all_paragraphs = document.paragraphs.copy()
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_paragraphs += cell.paragraphs

        for paragraph in all_paragraphs:
            if tag_text in paragraph.text:
                paragraph.clear()
                cursor = WordDocCursor(document, paragraph)
                break
 
        return cursor

    def new_bible_stream_handler(self, runner, cursor):
        return WordBibleStreamHandler(cursor)

    def text_found(self, document, search_text):
        '''Searches document for search_text, and returns True if found, else False.
        '''
        # Collect all the paragraphs from the document body plus any in table cells
        all_paragraphs = document.paragraphs.copy()
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_paragraphs += cell.paragraphs
                    
        for paragraph in all_paragraphs:
            if search_text in paragraph.text:
                return True
        
        # We only get to here if not found
        return False

    def insert_passage_group_tag(self, runner, document, cursor, is_first, tag_text):
        '''Called when expanding base template tags. Formats the document
        for a passage group header, and then actually inserts the tag_text.

        first_group is True if this is the first passage group header being inserted.
        '''
        if runner.plan.config.outputs[self.long_id].all_tables_insert_blank_paras and not is_first:
            cursor.add_new_para() # New paragraph for passage header

        cursor.current_para.style = get_style(document, Styles.PASSAGE_GROUP.value)
        cursor.add_text(tag_text)
        
        if runner.plan.config.outputs[self.long_id].all_tables_insert_blank_paras:
            cursor.add_new_para(copy_style=False) # Space para between passage header and text table

    def insert_passage_group_table(self, runner, document, cursor, is_first, table_text_array):
        '''Called when expanding base template tags. Inserts a passage group table at the current
        cursor point and then inserts the text in table_text_array into that table, adding any
        formatting necessary.
        '''
        table = cursor.add_new_table(len(table_text_array), len(table_text_array[0]),
                                     get_style(document, Styles.TEXT_TABLE_HORIZ.value))

        if is_first:
            # The first row will be the version name tags.
            table.show_header_row = True
            for column_index in range(len(table_text_array[0])):
                cell = table.cell(0, column_index)
                paragraph = cell.paragraphs[0]
                paragraph.style = get_style(document, Styles.VERSION_NAME.value)
                paragraph.text = table_text_array[0][column_index]
                # First col align-left. Last col alight right. Other cols align middle.
                if column_index == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif column_index == (len(runner.version_cols) - 1):
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # The last row (or only row, if there's just one row) contains the text tags.
        # Lines are separated by a '\n' character, which we need to convert into separate
        # paragraphs.
        for column_index in range(len(table_text_array[-1])):
            cell = table.cell(-1, column_index)
            paragraph = cell.paragraphs[0]
            text_lines = table_text_array[-1][column_index].split('\n') 
            for line_index in range(len(text_lines)):
                paragraph.style = get_style(document, Styles.PARAGRAPH_COL.value.format(column_symbols[column_index]))
                paragraph.text = text_lines[line_index]
                if line_index < (len(text_lines) - 1):
                    paragraph = cell.add_paragraph()

    def insert_copyright_table(self, runner, document, cursor, table_text_array):
        '''Called when expanding base template tags. If cursor is not None, inserts a copyright
        table at the current cursor point. If cursor is None, inserts a copyright table at the
        end of the document.
        '''
        table = None
        if cursor is not None:
            # We've already expanded an MS_ALL_TABLES tag, so insert the copyright table
            # at the current cursor location.
            table = cursor.add_new_table(1, len(table_text_array[0]),
                                            get_style(document, Styles.COPYRIGHT_TABLE_HORIZ.value))
        else:
            # We didn't expand an MS_ALL_TABLES tag, so just insert the copyright table at
            # the end of the document
            table = document.add_table(1, len(table_text_array[0]),
                                        get_style(document, Styles.COPYRIGHT_TABLE_HORIZ.value))
        # Expand the copyright table
        for column_index in range(len(table_text_array[0])):
            cell = table.cell(0, column_index)
            paragraph = cell.paragraphs[0]
            paragraph.style = get_style(document, Styles.COPYRIGHT.value)
            paragraph.add_run(table_text_array[0][column_index])

    def format_passage_tag(self, runner, document, cursor):
        '''Overridden from TaggedOuput. Performs any formatting needed prior to PASSAGE tag being inserted.
        '''
        cursor.current_para.style = get_style(document, Styles.PASSAGE.value)

    def format_bible_text_tag(self, runner, document, contents_index, column_symbol, bible_content, cursor):
        '''Overridden from TaggedOuput. Performs any formatting needed prior to Bible content being inserted.
        '''
        cursor.current_para.style = get_style(document, Styles.PARAGRAPH_COL.value.format(column_symbol))
        if runner.plan.config.outputs[self.long_id].apply_direct_formatting:
            cursor.run_font_family = bible_content.bible_version.font_family
            cursor.run_font_size = bible_content.bible_version.output_config[self.long_id].font_size
            cursor.para_is_bidi = bible_content.bible_version.is_rtl

    def format_text_join_tag(self, runner, document, column_symbol, bible_version, cursor):
        '''Overridden from TaggedOuput. Performs any formatting needed prior to join text being inserted.
        '''
        cursor.current_para.style = get_style(document, Styles.JOIN_COL.value.format(column_symbol))
        if runner.plan.config.outputs[self.long_id].apply_direct_formatting:
            cursor.run_font_family = bible_version.font_family
            cursor.run_font_size = bible_version.output_config[self.long_id].font_size
            cursor.para_is_bidi = bible_version.is_rtl

    def format_copyright_text_tag(self, runner, document, bible_content, cursor):
        '''Overridden from TaggedOuput. Performs any formatting needed prior to copyright text being inserted.
        '''
        cursor.current_para.style = get_style(document, Styles.COPYRIGHT.value)


class WordDocCursor(TaggedDocCursor):
    def __init__(self, document, current_para):
        super().__init__(document)
        self.current_para = current_para
        self.current_run = None

        self.run_font_family = None
        self.run_font_size = None
        self._para_is_bidi = False

        if self.current_run is None:
            self.add_new_run()

    @property
    def para_is_bidi(self) -> bool:
        return self._para_is_bidi
    
    @para_is_bidi.setter
    def para_is_bidi(self, value: bool):
        self._para_is_bidi = value
        set_para_formatting(self.current_para, value)

    def add_new_para(self, text=None, *, copy_style=True):
        prev_para = self.current_para
        new_p = OxmlElement('w:p')
        prev_para._p.addnext(new_p)
        self.current_para = Paragraph(new_p, prev_para._parent)

        if copy_style:
            self.current_para.style = prev_para.style

        set_para_formatting(self.current_para, self.para_is_bidi)

        #
        # TODO: We may need to copy some properties from the old paragraph to the new.
        #        
        self.add_new_run(text)

    def add_new_run(self, text=None):
        #
        # TODO: If this is not the first run in the paragraph, copy formatting from the previous run
        #       to the new run.
        #
        self.current_run = self.current_para.add_run()
        set_run_formatting(self.current_run, self.run_font_family, self.run_font_size, self.para_is_bidi)
        if text is not None:
            self.add_text(text)

    def add_text(self, text, *, convert_newlines=True):
        if convert_newlines:
            # Convert any newlines into new paragraphs
            text_lines = text.split('\n')
        else:
            text_lines = [text]
        for line_index in range(len(text_lines)):
            self.current_run.add_text(text_lines[line_index])
            if line_index < (len(text_lines) - 1):
                self.add_new_para(copy_style=True)

    def add_new_table(self, num_rows, num_cols, style=None, after_para=True):
        # Technique is from: 
        # https://github.com/python-openxml/python-docx/issues/156#issuecomment-77674193

        table = self.document.add_table(num_rows, num_cols, style)
        
        if after_para:
            # Add table after current paragraph. However, just moving the new table after
            # the current paragraph won't automatically update the current paragraph,
            # which will remain 'stuck' before the table. So instead we add a new blank
            # paragraph (which becomes the new current paragraph), and move the table
            # before it. That way we end up with a reference to the current paragraph
            # which is after the table, as we require.
            self.add_new_para(copy_style=False) # Post-table paragraph       
            self.current_para._p.addprevious(table._tbl)
        else:
            # Add table before current paragraph
            self.current_para._p.addprevious(table._tbl)

        # By default, turn off all special rows and columns
        table.show_header_row = False
        table.show_total_row = False
        table.show_header_column = False
        table.show_last_column = False
        table.show_banded_rows = False
        table.show_banded_columns = False

        return table
    

class WordBibleStreamHandler(OutputBibleStreamHandler):
    def __init__(self, word_doc_cursor):
        self.cursor = word_doc_cursor
        self.text_started = False       # True once any text has been added

    def add_text(self, text):
        self.cursor.add_text(text)
        if len(text.strip()) > 0:
            self.text_started = True

    def add_start_paragraph(self, is_poetry=False):
        # We only add a new paragraph if we've already added some text. This avoids an unnecessary blank
        # paragraph at the start.
        if self.text_started:
            self.cursor.add_new_para()
        if is_poetry:                
            self.cursor.current_para.paragraph_format.first_line_indent = docx.shared.Cm(-0.5)
            self.cursor.current_para.paragraph_format.left_indent = docx.shared.Cm(1)

    def add_end_paragraph(self):
        pass

    def add_line_break(self):
        self.cursor.current_run.add_break(docx.enum.text.WD_BREAK.LINE)

    def add_start_chap_num(self):
        self.cursor.add_new_run()
        self.cursor.current_run.font.bold = True

    def add_end_chap_num(self):
        self.cursor.add_new_run()

    def add_start_verse_num(self):
        self.cursor.add_new_run()
        self.cursor.current_run.font.bold = True
        self.cursor.current_run.font.superscript = True

    def add_end_verse_num(self):
        self.cursor.add_new_run()

    def add_start_small_caps(self):
        self.cursor.add_new_run()
        self.cursor.current_run.font.small_caps = True

    def add_end_small_caps(self):
        self.cursor.add_new_run()


class WordPlanConfig(OutputPlanConfig):
    def __init__(self, bible_output):
        super().__init__(bible_output)
        self.join_passage_text = "\n...\n"
        self.all_tables_insert_blank_paras = True
        self.apply_direct_formatting = False

    def new_config_widget(self):
        return WordPlanConfigPanel(None)


class WordVersionConfig(OutputVersionConfig):
    def __init__(self, bible_output):
        super().__init__(bible_output)
        self.font_size = 0

    def new_config_subform(self):
        return WordVersionConfigPanel(None)


class WordPlanRun(OutputPlanRun):
    def __init__(self, plan):
        super().__init__(plan)


#
# Module-level functions
#

def get_style(document, style_name):
    '''Look up a style by name in the document, and return style object if found, or None if not found.
    '''
    if style_name in document.styles:
        return document.styles[style_name]
    else:
        return None

def set_run_formatting(style_or_run: docx.styles.style.CharacterStyle | docx.text.run.Run,
                        font_family: str, font_size: str | float, is_rtl: bool) -> None:
    '''Manually set run formatting not yet available in the python-docx API.'''
    if font_family is not None and len(font_family) > 0:
        style_or_run.font.name = font_family
        # Ensures font name is also applied to East Asian fonts
        # See https://github.com/python-openxml/python-docx/issues/154#issuecomment-77707775
        rPr = style_or_run._element.rPr
        rPr.rFonts.set(qn('w:eastAsia'), font_family)
        # And we use the same technique to apply the font to other scripts as well:
        rPr.rFonts.set(qn('w:cs'), font_family) # Complex-scripts
        rPr.rFonts.set(qn('w:hAnsi'), font_family) # Any other scripts
    
    if font_size is not None and float(font_size) > 0:
        style_or_run.font.size = docx.shared.Pt(font_size)

    if is_rtl:
        style_or_run.font.rtl = True    # Needed to ensure right-to-left punctuation displays correctly

def set_para_formatting(style_or_para: docx.styles.style.ParagraphStyle | docx.text.paragraph.Paragraph,
                        is_bidi: bool) -> None:
    '''Manually set paragraph formatting not yet available in the python-docx API.'''
    if is_bidi:
        # Technique is from https://github.com/python-openxml/python-docx/issues/1411#issuecomment-2198293545
        pPr = style_or_para._element.get_or_add_pPr()
        bidi = OxmlElement("w:bidi")
        pPr.insert_element_before(bidi,
            *(
                "w:adjustRightInd",
                "w:snapToGrid",
                "w:spacing",
                "w:ind",
                "w:contextualSpacing",
                "w:mirrorIndents",
                "w:suppressOverlap",
                "w:jc",
                "w:textDirection",
                "w:textAlignment",
                "w:textboxTightWrap",
                "w:outlineLvl",
                "w:divId",
                "w:cnfStyle",
                "w:rPr",
                "w:sectPr",
                "w:pPrChange",
            )
        )

